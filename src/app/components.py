"""Reusable Streamlit components for each dashboard tab."""

from __future__ import annotations

import io
import logging

import pandas as pd
import streamlit as st

from src.config import COMPLETE_WEEKS, SCALE_FACTOR
from src.nlp_analysis import add_sentiment_columns
from src.visualizations import (
    plot_bpo_comparison,
    plot_category_treemap,
    plot_channel_distribution,
    plot_chatbot_escalation_by_category,
    plot_correlation_bar,
    plot_cost_efficiency_scatter,
    plot_csat_heatmap,
    plot_effort_impact_matrix,
    plot_frt_boxplot_by_team,
    plot_frustration_by_category,
    plot_heatmap_hourly,
    plot_kpi_trend,
    plot_multi_trend,
    plot_resolution_boxplot_by_team,
    plot_resolution_funnel,
    plot_sentiment_by_dimension,
    plot_sentiment_distribution,
    plot_team_comparison,
    plot_team_summary_table,
)

logger = logging.getLogger(__name__)


# ──────────────────────────────────────────────────────────────────────────────
# Shared helpers
# ──────────────────────────────────────────────────────────────────────────────

def _get_executive_insights() -> dict[str, str]:
    """Retrieve executive insights from session state (agent result)."""
    agent_res = st.session_state.get("agent_result")
    if agent_res and isinstance(agent_res.get("executive_insights"), dict):
        return agent_res["executive_insights"]
    return {}


def _render_insight_box(text: str, *, label: str = "Insight") -> None:
    """Render a styled insight callout if text is non-empty."""
    if not text or not text.strip():
        return
    st.info(f"**{label}:** {text}", icon="💡")


# ──────────────────────────────────────────────────────────────────────────────
# Tab 1 — Dashboard Overview
# ──────────────────────────────────────────────────────────────────────────────
def render_dashboard_tab(df: pd.DataFrame, analytics: dict) -> None:
    """Render the main KPI dashboard tab."""
    kpi = analytics["kpi"]
    wow = analytics.get("wow", {})
    deltas = wow.get("deltas", {})
    cur_wk = wow.get("current_week")
    pri_wk = wow.get("prior_week")
    date_ranges = analytics.get("week_date_ranges", {})

    if cur_wk and pri_wk:
        cur_dates = date_ranges.get(cur_wk, "")
        pri_dates = date_ranges.get(pri_wk, "")
        _cw = analytics.get("complete_weeks", COMPLETE_WEEKS)
        st.caption(
            f"Overall sample ({kpi['total_tickets']:,} tickets, Weeks {min(_cw)}-{max(_cw)}) "
            f"| Week-over-week: Week {cur_wk} ({cur_dates}) vs Week {pri_wk} ({pri_dates})"
        )

    def _delta_str(key: str, fmt: str = "abs", prefix: str = "", suffix: str = "") -> str | None:
        """Format a WoW delta for st.metric display."""
        d = deltas.get(key)
        if not d:
            return None
        v = d["abs_change"]
        if fmt == "pct":
            return f"{v:+.1%}{suffix}"
        if fmt == "money":
            return f"{prefix}{v:+.2f}{suffix}"
        if fmt == "int":
            return f"{v:+.0f}{suffix}"
        return f"{v:+.1f}{suffix}"

    # Primary KPI row — native st.metric cards with WoW deltas
    cols = st.columns(5)
    cols[0].metric("Total Tickets", f"{kpi['total_tickets']:,}",
                   delta=_delta_str("total_tickets", fmt="int"))
    cols[1].metric("Resolution Rate", f"{kpi['resolution_rate']:.1%}",
                   delta=_delta_str("resolution_rate", fmt="pct"))
    cols[2].metric("Avg FRT", f"{kpi['avg_first_response_min']:.0f} min",
                   delta=_delta_str("avg_first_response_min", suffix=" min"),
                   delta_color="inverse")
    cols[3].metric("Avg CSAT", f"{kpi['avg_csat']:.2f} / 5",
                   delta=_delta_str("avg_csat", fmt="money"))
    cols[4].metric("Avg Cost", f"${kpi['avg_cost_usd']:.2f}",
                   delta=_delta_str("avg_cost_usd", fmt="money", prefix="$"),
                   delta_color="inverse")

    # Secondary KPI row
    cols2 = st.columns(4)
    cols2[0].metric("Escalation Rate", f"{kpi['escalation_rate']:.1%}",
                    delta=_delta_str("escalation_rate", fmt="pct"),
                    delta_color="inverse")
    cols2[1].metric("Abandonment Rate", f"{kpi['abandonment_rate']:.1%}",
                    delta=_delta_str("abandonment_rate", fmt="pct"),
                    delta_color="inverse")
    cols2[2].metric("Contacts / Ticket", f"{kpi['avg_contacts_per_ticket']:.1f}",
                    delta=_delta_str("avg_contacts_per_ticket"),
                    delta_color="inverse")
    cols2[3].metric(
        "Total Cost (Weekly)",
        f"${kpi['total_cost_usd']:,.0f}",
        delta=_delta_str("total_cost_usd", fmt="money", prefix="$"),
        delta_color="inverse",
    )


    # ── Data Quality Summary ──────────────────────────────────────────────
    quality_report = st.session_state.get("quality_report", {})
    cleaning_log = st.session_state.get("cleaning_log", {})
    if quality_report or cleaning_log:
        with st.expander("Data Quality Pipeline", expanded=False):
            dq1, dq2 = st.columns(2)
            raw_completeness = quality_report.get("completeness_score", 0)
            # After cleaning + imputation the dataset is nearly complete
            imputed_count = sum(
                v for k, v in cleaning_log.items() if k.endswith("_imputed")
            )
            total_cells = quality_report.get("total_rows", 0) * quality_report.get("total_columns", 0)
            final_completeness = (
                (raw_completeness * total_cells + imputed_count) / max(total_cells, 1)
                if total_cells > 0 else raw_completeness
            )
            dq1.metric("Raw Completeness", f"{raw_completeness:.1%}")
            dq2.metric("Post-Pipeline Completeness", f"{min(final_completeness, 1.0):.1%}")

            st.caption("Cleaning steps applied:")
            log_items = []
            for k, v in cleaning_log.items():
                label = k.replace("_", " ").title()
                log_items.append(f"- **{label}**: {v:,} rows")
            st.markdown("\n".join(log_items))


    # Charts Row 1: Team + Channel
    col1, col2 = st.columns(2)
    with col1:
        team_df = pd.DataFrame(analytics["teams"])
        fig = plot_team_comparison(
            team_df, "resolution_rate", "Resolution Rate by Team"
        )
        st.plotly_chart(fig, width="stretch")
    with col2:
        fig = plot_channel_distribution(df)
        st.plotly_chart(fig, width="stretch")

    # Charts Row 2: Funnel + Heatmap
    col1, col2 = st.columns(2)
    with col1:
        fig = plot_resolution_funnel(df)
        st.plotly_chart(fig, width="stretch")
    with col2:
        fig = plot_heatmap_hourly(df)
        st.plotly_chart(fig, width="stretch")

    # Charts Row 3: Cost Efficiency + Category Treemap
    col1, col2 = st.columns(2)
    with col1:
        team_df = pd.DataFrame(analytics["teams"])
        fig = plot_cost_efficiency_scatter(team_df)
        st.plotly_chart(fig, width="stretch")
    with col2:
        fig = plot_category_treemap(df)
        st.plotly_chart(fig, width="stretch")

    # Charts Row 4: FRT + Resolution box plots
    st.divider()
    col1, col2 = st.columns(2)
    with col1:
        fig = plot_frt_boxplot_by_team(df)
        st.plotly_chart(fig, width="stretch")
    with col2:
        fig = plot_resolution_boxplot_by_team(df)
        st.plotly_chart(fig, width="stretch")

    # Charts Row 5: CSAT Heatmap + BPO Comparison
    col1, col2 = st.columns(2)
    with col1:
        fig = plot_csat_heatmap(df)
        st.plotly_chart(fig, width="stretch")
    with col2:
        fig = plot_bpo_comparison(analytics["teams"])
        st.plotly_chart(fig, width="stretch")

    # ── Chatbot Escalation Section ────────────────────────────────────────
    chatbot = analytics.get("chatbot", {})
    if chatbot:
        st.divider()
        st.subheader("AI Chatbot Performance")
        cb1, cb2, cb3 = st.columns(3)
        cb1.metric(
            "Chatbot Tickets",
            f"{chatbot.get('total_chatbot_tickets', 0):,}",
        )
        cb2.metric(
            "Escalation Rate",
            f"{chatbot.get('overall_escalation_rate', 0):.1%}",
        )
        cb3.metric(
            "Escalated Tickets",
            f"{chatbot.get('total_escalated', 0):,}",
        )
        fig = plot_chatbot_escalation_by_category(chatbot)
        st.plotly_chart(fig, width="stretch")


# ──────────────────────────────────────────────────────────────────────────────
# Tab 2 — Opportunities
# ──────────────────────────────────────────────────────────────────────────────
def render_opportunities_tab(
    df: pd.DataFrame, analytics: dict, nlp_results: dict,
    agent_result: dict | None = None,
) -> None:
    """Render agent-identified improvement opportunities."""
    # Agent-generated intro paragraph — rendered above the heading
    insights = _get_executive_insights()
    intro_text = insights.get("opportunities_intro", "")
    if intro_text:
        st.markdown(
            f'<div style="background: #f0f7eb; border-left: 4px solid #53A318; '
            f'padding: 12px 16px; border-radius: 4px; margin-bottom: 1rem; '
            f'font-size: 0.95em; color: #3d3a2a;">{intro_text}</div>',
            unsafe_allow_html=True,
        )

    st.subheader("Top Improvement Opportunities")
    st.caption(
        "AI-identified improvement opportunities from deterministic analytics & agent reasoning"
    )

    # Use agent-generated opportunities (not hardcoded)
    if agent_result is None:
        agent_result = st.session_state.get("agent_result")
    opportunities: list[dict] = []
    if agent_result and isinstance(agent_result.get("opportunities"), list):
        opportunities = agent_result["opportunities"]

    # Deterministic fallback when agent opportunities are unavailable
    using_fallback = False
    if not opportunities:
        opportunities = _compute_deterministic_opportunities(analytics, nlp_results)
        if opportunities:
            using_fallback = True
            st.info(
                "Showing **data-driven opportunities** computed deterministically. "
                "AI-generated insights were unavailable (LLM rate-limit or error).",
                icon="📊",
            )
        else:
            st.warning(
                "No improvement opportunities available. "
                "Try re-running the analysis."
            )
            return

    st.divider()

    for i, opp in enumerate(opportunities):
        rank = opp.get("rank", i + 1)
        title = opp.get("title", f"Opportunity {rank}")
        impact = opp.get("impact_estimate", "N/A")
        effort = str(opp.get("effort", "N/A")).title()
        priority = opp.get("priority_score", 0)
        category = str(opp.get("category", "N/A")).title()

        if isinstance(priority, (int, float)):
            p_label = (
                "CRITICAL" if priority >= 80
                else ("HIGH" if priority >= 60 else "MEDIUM")
            )
        else:
            p_label = str(priority)

        with st.expander(
            f"{title}  |  {impact}  |  {p_label}",
            expanded=(i == 0),
        ):
            col1, col2, col3 = st.columns(3)
            _metric_card = (
                '<div style="border-left: 3px solid #53A318; padding: 8px 12px;">'
                '<span style="color: #53A318; font-size: 0.85em; font-weight: 600;">{label}</span><br>'
                '<span style="font-size: 1.3em; font-weight: 700; word-wrap: break-word;">{value}</span></div>'
            )
            col1.markdown(_metric_card.format(label="IMPACT", value=impact), unsafe_allow_html=True)
            col2.markdown(_metric_card.format(label="EFFORT", value=effort), unsafe_allow_html=True)
            col3.markdown(_metric_card.format(label="CATEGORY", value=category), unsafe_allow_html=True)

            desc = opp.get("description", "")
            if desc:
                st.markdown(f"**Analysis:** {desc}")

            rec = opp.get("recommendation", "")
            if rec:
                st.markdown("**Recommendation:**")
                if isinstance(rec, list):
                    for item in rec:
                        st.markdown(f"- {item}")
                elif isinstance(rec, str) and rec.startswith("["):
                    # Handle string representation of a list (e.g. "['item1', 'item2']")
                    import ast
                    try:
                        parsed = ast.literal_eval(rec)
                        if isinstance(parsed, list):
                            for item in parsed:
                                st.markdown(f"- {item}")
                        else:
                            st.markdown(rec)
                    except (ValueError, SyntaxError):
                        st.markdown(rec)
                elif "\n" in rec:
                    # Multi-line text: render each line as a bullet
                    for line in rec.strip().split("\n"):
                        line = line.strip().lstrip("- ").lstrip("•").strip()
                        if line:
                            st.markdown(f"- {line}")
                else:
                    # Single string, may have numbered steps like "1. ... 2. ... 3. ..."
                    import re as _re
                    steps = _re.split(r"\s*\d+\.\s+", rec)
                    steps = [s.strip() for s in steps if s.strip()]
                    if len(steps) > 1:
                        for step in steps:
                            st.markdown(f"- {step}")
                    else:
                        st.markdown(rec)

            supporting = opp.get("supporting_data")
            if supporting:
                st.markdown("**Supporting Data:**")
                if isinstance(supporting, list):
                    for point in supporting:
                        st.markdown(f"- {point}")
                elif isinstance(supporting, dict):
                    for k, v in supporting.items():
                        st.markdown(f"- **{k}:** {v}")
                else:
                    st.markdown(f"- {supporting}")

    # Effort–Impact Matrix
    if opportunities:
        st.divider()
        st.subheader("Effort–Impact Matrix")
        fig = plot_effort_impact_matrix(opportunities)
        st.plotly_chart(fig, width="stretch")


def _compute_deterministic_opportunities(
    analytics: dict, nlp_results: dict,
) -> list[dict]:
    """Compute top 5 opportunities deterministically from analytics data.

    Uses the same formulas as the EDA notebook's opportunity sizing section.
    This serves as a fallback when the LLM agent is unavailable.
    """
    opps: list[dict] = []
    kpi = analytics.get("kpi", {})
    chatbot = analytics.get("chatbot", {})
    teams_list = analytics.get("teams", [])
    channels_list = analytics.get("channels", [])

    teams_df = pd.DataFrame(teams_list) if teams_list else pd.DataFrame()

    # Opp 1: Chatbot escalation reduction
    if chatbot:
        total_bot = chatbot.get("total_chatbot_tickets", 0)
        esc_rate = chatbot.get("overall_escalation_rate", 0)
        current_esc = int(total_bot * esc_rate)
        target_rate = 0.15
        avoided = max(0, current_esc - int(total_bot * target_rate))
        # avg cost of an escalated ticket (use overall avg as proxy)
        avg_cost = kpi.get("avg_cost_usd", 4.0)
        sample_savings = avoided * avg_cost
        annual_savings = sample_savings * SCALE_FACTOR
        opps.append({
            "rank": 1,
            "title": "Reduce Chatbot Escalation Rate",
            "impact_estimate": f"~${annual_savings:,.0f}/year",
            "effort": "Medium",
            "priority_score": 90,
            "category": "Automation",
            "description": (
                f"The AI chatbot escalates {esc_rate:.1%} of its {total_bot:,} tickets. "
                f"Reducing to {target_rate:.0%} would avoid ~{avoided:,} escalations per sample period. "
                "Root causes: merchant_issue and billing categories drive the highest chatbot escalation rates."
            ),
            "recommendation": (
                "1. Retrain chatbot on billing and merchant_issue categories with better resolution flows. "
                "2. Add pre-escalation confirmation step. "
                "3. Route complex categories directly to human agents."
            ),
            "supporting_data": {
                "current_escalation_rate": f"{esc_rate:.1%}",
                "target_rate": f"{target_rate:.0%}",
                "chatbot_tickets": f"{total_bot:,}",
                "avoided_escalations": f"{avoided:,}",
            },
        })

    # Opp 2: BPO Vendor B performance gap
    if not teams_df.empty:
        vendor_b = teams_df[teams_df["team"] == "bpo_vendorB"]
        vendor_a = teams_df[teams_df["team"] == "bpo_vendorA"]
        if not vendor_b.empty and not vendor_a.empty:
            vb_csat = vendor_b.iloc[0].get("avg_csat", 0)
            va_csat = vendor_a.iloc[0].get("avg_csat", 0)
            vb_esc = vendor_b.iloc[0].get("escalation_rate", 0)
            vb_tickets = vendor_b.iloc[0].get("ticket_count", 0)
            csat_gap = va_csat - vb_csat
            opps.append({
                "rank": 2,
                "title": "Close BPO Vendor B Performance Gap",
                "impact_estimate": f"CSAT +{csat_gap:.2f} for {vb_tickets:,} tickets",
                "effort": "Medium",
                "priority_score": 82,
                "category": "Customer Satisfaction",
                "description": (
                    f"Vendor B CSAT ({vb_csat:.2f}) is {csat_gap:.2f} points below Vendor A ({va_csat:.2f}). "
                    f"Vendor B escalation rate is {vb_esc:.1%}. "
                    "This affects a significant volume of tickets and directly impacts customer satisfaction."
                ),
                "recommendation": (
                    "1. Conduct root-cause analysis of Vendor B's lowest-CSAT categories. "
                    "2. Implement targeted training program on top 3 underperforming categories. "
                    "3. Set CSAT improvement SLA in vendor contract."
                ),
                "supporting_data": {
                    "vendor_b_csat": f"{vb_csat:.2f}",
                    "vendor_a_csat": f"{va_csat:.2f}",
                    "csat_gap": f"{csat_gap:.2f}",
                    "vendor_b_escalation_rate": f"{vb_esc:.1%}",
                },
            })

    # Opp 3: Reduce abandonment
    aband_rate = kpi.get("abandonment_rate", 0)
    total_tickets = kpi.get("total_tickets", 0)
    abandoned_count = int(total_tickets * aband_rate)
    clv_lost = abandoned_count * 15  # $15 CLV assumption
    annual_clv = clv_lost * SCALE_FACTOR
    opps.append({
        "rank": 3,
        "title": "Reduce Ticket Abandonment Rate",
        "impact_estimate": f"~${annual_clv:,.0f}/year in retained CLV",
        "effort": "Medium",
        "priority_score": 78,
        "category": "Customer Satisfaction",
        "description": (
            f"Current abandonment rate is {aband_rate:.1%} ({abandoned_count:,} tickets). "
            "Each abandoned ticket represents a customer who gave up. "
            "At an estimated $15 CLV per lost customer, this is significant revenue at risk."
        ),
        "recommendation": (
            "1. Correlate abandonment with FRT — reduce wait times in high-abandonment channels. "
            "2. Add proactive follow-up for tickets idle > 30 min. "
            "3. Implement queue position / ETA notifications."
        ),
        "supporting_data": {
            "abandonment_rate": f"{aband_rate:.1%}",
            "abandoned_tickets": f"{abandoned_count:,}",
            "estimated_clv_per_customer": "$15",
        },
    })

    # Opp 4: Email channel optimization
    email_data = None
    for ch in channels_list:
        if ch.get("channel") == "email":
            email_data = ch
            break
    if email_data:
        email_frt = email_data.get("avg_frt_min", 0)
        email_cost = email_data.get("avg_cost_usd", 0)
        email_count = email_data.get("ticket_count", 0)
        opps.append({
            "rank": 4,
            "title": "Optimize Email Channel Efficiency",
            "impact_estimate": f"Improve FRT from {email_frt:.0f} min",
            "effort": "Low",
            "priority_score": 70,
            "category": "Process Efficiency",
            "description": (
                f"Email handles {email_count:,} tickets with avg FRT of {email_frt:.0f} min "
                f"and cost of ${email_cost:.2f}/ticket. "
                "Many email tickets could be deflected to chat or self-service for faster resolution."
            ),
            "recommendation": (
                "1. Implement auto-reply with top FAQ answers + self-service links. "
                "2. Add auto-classification to route simple queries to chatbot. "
                "3. Set FRT SLA target of 30 min for email."
            ),
            "supporting_data": {
                "email_volume": f"{email_count:,}",
                "email_avg_frt": f"{email_frt:.0f} min",
                "email_avg_cost": f"${email_cost:.2f}",
            },
        })

    # Opp 5: Improve CSAT collection rate
    csat_collection = kpi.get("csat_collection_rate", 0)
    opps.append({
        "rank": 5,
        "title": "Increase CSAT Survey Collection Rate",
        "impact_estimate": "Improve data quality for decision-making",
        "effort": "Low",
        "priority_score": 60,
        "category": "Process Efficiency",
        "description": (
            f"CSAT collection rate is {csat_collection:.1%}. "
            "Missing CSAT data reduces confidence in satisfaction metrics and hides problem areas. "
            "Chatbot-handled tickets have the lowest collection rates."
        ),
        "recommendation": (
            "1. Add mandatory CSAT prompt at end of every chat/phone interaction. "
            "2. Send automated survey email within 1 hour of ticket resolution. "
            "3. Target 90%+ collection rate."
        ),
        "supporting_data": {
            "current_collection_rate": f"{csat_collection:.1%}",
            "target_rate": "90%+",
        },
    })

    return opps[:5]


# ──────────────────────────────────────────────────────────────────────────────
# Tab 3 — Trends
# ──────────────────────────────────────────────────────────────────────────────
def render_trends_tab(df: pd.DataFrame, analytics: dict) -> None:
    """Render the weekly trends analysis tab."""

    # Week context header
    wow = analytics.get("wow", {})
    cur_wk = wow.get("current_week")
    pri_wk = wow.get("prior_week")
    all_weeks = sorted(analytics.get("complete_weeks", COMPLETE_WEEKS))
    date_ranges = analytics.get("week_date_ranges", {})

    # Build date-aware week labels
    cur_dates = date_ranges.get(cur_wk, "")
    pri_dates = date_ranges.get(pri_wk, "")
    first_dates = date_ranges.get(min(all_weeks), "") if all_weeks else ""
    last_dates = date_ranges.get(max(all_weeks), "") if all_weeks else ""

    if all_weeks:
        st.markdown(
            f"**Data spans Weeks {min(all_weeks)}-{max(all_weeks)}** "
            f"({first_dates} to {last_dates}, {len(all_weeks)} complete weeks). "
            f"Week-over-week comparison: **Week {cur_wk}** ({cur_dates}) "
            f"vs **Week {pri_wk}** ({pri_dates})."
        )
    else:
        st.warning("No complete weeks detected in the data.")

    # WoW summary cards — show actual Week N values + delta from prior week
    deltas = wow.get("deltas", {})
    cur_kpis = wow.get("current", {})
    if deltas and cur_kpis:
        st.caption(f"Week {cur_wk} ({cur_dates}) values — change vs Week {pri_wk} ({pri_dates})")
        cols = st.columns(5)
        _wow_metrics = [
            ("Resolution Rate", "resolution_rate", "pct"),
            ("Escalation Rate", "escalation_rate", "pct_inv"),
            ("Avg CSAT", "avg_csat", "f2"),
            ("Avg FRT (min)", "avg_first_response_min", "f1_inv"),
            ("Avg Cost", "avg_cost_usd", "money_inv"),
        ]
        for col, (label, key, fmt) in zip(cols, _wow_metrics):
            d = deltas.get(key, {})
            v = d.get("abs_change", 0)
            cur_val = cur_kpis.get(key)
            if cur_val is None:
                continue
            # Format the primary value
            if fmt == "pct":
                val_str = f"{cur_val:.1%}"
                delta_str = f"{v:+.1%}"
                col.metric(label, val_str, delta=delta_str)
            elif fmt == "pct_inv":
                val_str = f"{cur_val:.1%}"
                delta_str = f"{v:+.1%}"
                col.metric(label, val_str, delta=delta_str, delta_color="inverse")
            elif fmt == "f2":
                val_str = f"{cur_val:.2f}"
                delta_str = f"{v:+.2f}"
                col.metric(label, val_str, delta=delta_str)
            elif fmt == "f1_inv":
                val_str = f"{cur_val:.0f} min"
                delta_str = f"{v:+.1f} min"
                col.metric(label, val_str, delta=delta_str, delta_color="inverse")
            elif fmt == "money_inv":
                val_str = f"${cur_val:.2f}"
                delta_str = f"${v:+.2f}"
                col.metric(label, val_str, delta=delta_str, delta_color="inverse")

    # Executive insights for overall trends
    insights = _get_executive_insights()
    _render_insight_box(insights.get("trends_insight", ""), label="Trends Insight")


    # Multi-metric overview chart (show several KPIs at once)
    derived_df = pd.DataFrame(analytics["derived_trends"])
    if not derived_df.empty:
        first_wk_dates = date_ranges.get(all_weeks[0] if all_weeks else 0, "")
        last_wk_dates = date_ranges.get(all_weeks[-1] if all_weeks else 0, "")
        st.subheader(f"Key Metrics Over Time ({first_wk_dates} to {last_wk_dates})")
        fig = plot_multi_trend(derived_df, week_date_ranges=date_ranges)
        st.plotly_chart(fig, width="stretch")


    # Deep-dive: pick individual raw metrics
    raw_df = pd.DataFrame(analytics["raw_trends"])
    if not raw_df.empty:
        st.subheader("Metric Deep Dive")
        raw_metric = st.selectbox(
            "Select metric",
            [
                "csat_score",
                "first_response_min",
                "resolution_min",
                "cost_usd",
                "contacts_per_ticket",
            ],
            format_func=lambda x: x.replace("_", " ").title(),
        )
        fig = plot_kpi_trend(
            raw_df,
            raw_metric,
            f"{raw_metric.replace('_', ' ').title()} — Weekly Average",
            week_date_ranges=date_ranges,
        )
        st.plotly_chart(fig, width="stretch")
        # Agent insight for metric trends
        _render_insight_box(
            _get_executive_insights().get("trends_insight", ""),
            label="Metric Trend Insight",
        )

    st.divider()

    # Team performance table (compact, replaces big bar chart)
    st.subheader("Team Performance Summary")
    team_df = pd.DataFrame(analytics["teams"])
    fig = plot_team_summary_table(team_df)
    st.plotly_chart(fig, width="stretch")
    # Agent insight for team performance
    _render_insight_box(
        _get_executive_insights().get("team_performance_insight", ""),
        label="Team Performance Insight",
    )

    st.divider()

    # CSAT Correlation analysis (real chart, not empty dataframe)
    if analytics.get("correlations"):
        st.subheader("CSAT Correlation Analysis")
        st.caption("Pearson correlation between CSAT and operational metrics")
        fig = plot_correlation_bar(analytics["correlations"])
        st.plotly_chart(fig, width="stretch")

        # Executive insight for correlations
        insights = _get_executive_insights()
        _render_insight_box(insights.get("correlation_insight", ""), label="Correlation Insight")


# ──────────────────────────────────────────────────────────────────────────────
# Tab 4 — NLP Insights
# ──────────────────────────────────────────────────────────────────────────────
def render_nlp_tab(df: pd.DataFrame, nlp_results: dict) -> None:
    """Render the NLP insights tab."""

    if not nlp_results:
        st.warning("NLP analysis not available.")
        return

    # Sentiment overview
    col1, col2, col3 = st.columns(3)
    col1.metric("Avg Sentiment", f"{nlp_results.get('avg_sentiment_polarity', 0):.3f}")
    col2.metric("Frustration Rate", f"{nlp_results.get('frustration_rate', 0):.1%}")
    col3.metric(
        "Frustrated Tickets",
        f"{nlp_results.get('total_frustrated_tickets', 0):,}",
    )

    # Sentiment distribution + Frustration side by side
    col1, col2 = st.columns(2)
    with col1:
        if nlp_results.get("sentiment_distribution"):
            # st.subheader("Sentiment Distribution")
            df_with_sent = add_sentiment_columns(df)
            fig = plot_sentiment_distribution(df_with_sent)
            st.plotly_chart(fig, width="stretch")

    with col2:
        if nlp_results.get("frustration_by_category"):
            # st.subheader("Frustration Rate by Category")
            fig = plot_frustration_by_category(nlp_results["frustration_by_category"])
            st.plotly_chart(fig, width="stretch")

    # Executive insights for sentiment and frustration
    insights = _get_executive_insights()
    c1, c2 = st.columns(2)
    with c1:
        _render_insight_box(insights.get("sentiment_insight", ""), label="Sentiment Insight")
    with c2:
        _render_insight_box(insights.get("frustration_insight", ""), label="Frustration Insight")

    st.divider()

    # Topics — presented as a clean table
    if nlp_results.get("topics"):
        st.subheader("Detected Topics (TF-IDF Clustering)")
        topic_data = []
        for topic in nlp_results["topics"]:
            topic_data.append({
                "Topic": topic.get("label", f"Topic {topic['topic_id']}"),
                "Tickets": topic.get("count", 0),
                "Key Terms": ", ".join(topic["top_words"][:7]),
            })
        st.dataframe(
            pd.DataFrame(topic_data),
            hide_index=True,
        )
        _render_insight_box(
            _get_executive_insights().get("topic_insight", ""),
            label="Topic Insight",
        )

    # Sample frustrated messages
    if nlp_results.get("sample_frustrated_messages"):
        st.subheader("Sample Frustrated Messages")
        for msg in nlp_results["sample_frustrated_messages"][:5]:
            st.markdown(f'> *"{msg}"*')

    st.divider()

    # ── Cross-dimensional sentiment breakdowns ────────────────────────────
    st.subheader("Sentiment by Dimension")

    sent_tabs = st.tabs(["By Team", "By Channel", "By Priority"])

    with sent_tabs[0]:
        sent_by_team = nlp_results.get("sentiment_by_team", {})
        if sent_by_team:
            fig = plot_sentiment_by_dimension(sent_by_team, "Team")
            st.plotly_chart(fig, width="stretch")
        else:
            st.caption("No per-team sentiment data available.")

    with sent_tabs[1]:
        sent_by_channel = nlp_results.get("sentiment_by_channel", {})
        if sent_by_channel:
            fig = plot_sentiment_by_dimension(sent_by_channel, "Channel")
            st.plotly_chart(fig, width="stretch")
        else:
            st.caption("No per-channel sentiment data available.")

    with sent_tabs[2]:
        sent_by_priority = nlp_results.get("sentiment_by_priority", {})
        if sent_by_priority:
            fig = plot_sentiment_by_dimension(sent_by_priority, "Priority")
            st.plotly_chart(fig, width="stretch")
        else:
            st.caption("No per-priority sentiment data available.")

    # Agent insight for cross-dimensional sentiment
    _render_insight_box(
        _get_executive_insights().get("sentiment_dimension_insight", ""),
        label="Sentiment by Dimension Insight",
    )

    # Frustration ↔ CSAT correlation
    csat_by_frust = nlp_results.get("csat_by_frustration", {})
    if csat_by_frust:
        st.divider()
        st.subheader("Frustration Impact on CSAT")
        fc1, fc2 = st.columns(2)
        fc1.metric(
            "Frustrated Avg CSAT",
            f"{csat_by_frust.get('frustrated_avg_csat', 0):.2f}",
        )
        fc2.metric(
            "Non-Frustrated Avg CSAT",
            f"{csat_by_frust.get('non_frustrated_avg_csat', 0):.2f}",
        )
        gap = (
            csat_by_frust.get("non_frustrated_avg_csat", 0)
            - csat_by_frust.get("frustrated_avg_csat", 0)
        )
        if gap > 0:
            st.info(
                f"Frustrated customers rate **{gap:.2f} points lower** on CSAT. "
                "Reducing frustration drivers directly improves satisfaction scores.",
                icon="📉",
            )
        # Agent insight for frustration impact
        _render_insight_box(
            _get_executive_insights().get("frustration_insight", ""),
            label="Frustration Impact Insight",
        )


# ──────────────────────────────────────────────────────────────────────────────
# Tab 5 — Weekly Brief
# ──────────────────────────────────────────────────────────────────────────────

# ── Export helpers ────────────────────────────────────────────────────────────

def _generate_markdown_brief(
    analytics: dict,
    nlp_results: dict,
    quality_report: dict,
    cleaning_log: dict,
) -> str:
    """Build a deterministic KPI reference table as Markdown.

    Contains only grounded, data-driven numbers (KPI comparison table
    and data quality notes).  Narrative interpretation comes from the
    LLM agent's ``report_markdown`` output.
    """
    kpi = analytics["kpi"]
    chatbot = analytics["chatbot"]
    teams_list = analytics["teams"]
    best_team = max(teams_list, key=lambda t: t.get("avg_csat", 0))
    worst_team = min(teams_list, key=lambda t: t.get("avg_csat", 5))

    # Week-over-week deltas
    wow = analytics.get("wow", {})
    deltas = wow.get("deltas", {})
    _cw_brief = analytics.get("complete_weeks", COMPLETE_WEEKS)
    cur_wk = wow.get("current_week", max(_cw_brief) if _cw_brief else 0)
    pri_wk = wow.get("prior_week", (max(_cw_brief) - 1) if _cw_brief else 0)
    cur_kpis = wow.get("current", {})
    pri_kpis = wow.get("prior", {})

    # Date ranges for human-readable labels
    date_ranges = analytics.get("week_date_ranges", {})
    cur_dates = date_ranges.get(cur_wk, "")
    pri_dates = date_ranges.get(pri_wk, "")
    first_dates = date_ranges.get(min(_cw_brief), "") if _cw_brief else ""
    last_dates = date_ranges.get(max(_cw_brief), "") if _cw_brief else ""

    def _fmt_val(key: str, fmt: str = "pct") -> str:
        """Format a weekly KPI value for the table."""
        v = cur_kpis.get(key)
        if v is None:
            return "--"
        if fmt == "pct":
            return f"{v:.1%}"
        if fmt == "money":
            return f"${v:.2f}"
        if fmt == "time":
            return f"{v:.0f} min"
        if fmt == "int":
            return f"{v:,}"
        return f"{v:.2f}"

    def _fmt_prior(key: str, fmt: str = "pct") -> str:
        """Format a prior-week KPI value for the table."""
        v = pri_kpis.get(key)
        if v is None:
            return "--"
        if fmt == "pct":
            return f"{v:.1%}"
        if fmt == "money":
            return f"${v:.2f}"
        if fmt == "time":
            return f"{v:.0f} min"
        if fmt == "int":
            return f"{v:,}"
        return f"{v:.2f}"

    def _wow_cell(key: str, fmt: str = "pct") -> str:
        """Format a WoW delta for the Markdown table."""
        d = deltas.get(key)
        if not d:
            return "--"
        v = d["abs_change"]
        direction = d["direction"]
        arrow = "Improving" if direction == "improving" else ("Declining" if direction == "worsening" else "Flat")
        if fmt == "pct":
            return f"{v:+.1%} ({arrow})"
        if fmt == "money":
            return f"${v:+.2f} ({arrow})"
        if fmt == "time":
            return f"{v:+.1f} min ({arrow})"
        if fmt == "int":
            return f"{v:+,} ({arrow})"
        return f"{v:+.2f} ({arrow})"

    def _status(key: str) -> str:
        d = deltas.get(key, {})
        direction = d.get("direction", "stable")
        if direction == "improving":
            return "On Track"
        if direction == "worsening":
            return "At Risk"
        return "Stable"

    frustration_rate = nlp_results.get("frustration_rate", 0)
    avg_polarity = nlp_results.get("avg_sentiment_polarity", 0)

    # Weekly ticket counts for the executive summary
    cur_tickets = cur_kpis.get('total_tickets', 0)
    pri_tickets = pri_kpis.get('total_tickets', 0)
    ticket_delta = cur_tickets - pri_tickets
    ticket_dir = "up" if ticket_delta > 0 else ("down" if ticket_delta < 0 else "flat")

    # Compute weekly/monthly total cost estimates from the data as-is
    weekly_cost = cur_kpis.get('total_cost_usd', 0)
    monthly_cost_est = weekly_cost * 4  # ~4 weeks per month

    # Build the deterministic KPI reference (always accurate)
    data_section = f"""### Deterministic KPI Reference

> Data spans Weeks {min(_cw_brief)}-{max(_cw_brief)} ({first_dates} to {last_dates}, {len(_cw_brief)} complete weeks). Below compares the two most recent complete weeks.

---

#### KPI Dashboard  (Week {cur_wk} [{cur_dates}] vs Week {pri_wk} [{pri_dates}])

| Metric | Wk {cur_wk} | Wk {pri_wk} | WoW Change | Monthly Est. | Status |
|:-------|------:|------:|:-----------|------:|:-------|
| Total Tickets | {_fmt_val('total_tickets', 'int')} | {_fmt_prior('total_tickets', 'int')} | {_wow_cell('total_tickets', 'int')} | ~{cur_tickets * 4:,}/mo | -- |
| Resolution Rate | {_fmt_val('resolution_rate')} | {_fmt_prior('resolution_rate')} | {_wow_cell('resolution_rate')} | -- | {_status('resolution_rate')} |
| Escalation Rate | {_fmt_val('escalation_rate')} | {_fmt_prior('escalation_rate')} | {_wow_cell('escalation_rate')} | -- | {_status('escalation_rate')} |
| Abandonment Rate | {_fmt_val('abandonment_rate')} | {_fmt_prior('abandonment_rate')} | {_wow_cell('abandonment_rate')} | -- | {_status('abandonment_rate')} |
| Avg First Response | {_fmt_val('avg_first_response_min', 'time')} | {_fmt_prior('avg_first_response_min', 'time')} | {_wow_cell('avg_first_response_min', 'time')} | -- | {_status('avg_first_response_min')} |
| Avg Resolution Time | {_fmt_val('avg_resolution_min', 'time')} | {_fmt_prior('avg_resolution_min', 'time')} | {_wow_cell('avg_resolution_min', 'time')} | -- | {_status('avg_resolution_min')} |
| Avg CSAT | {_fmt_val('avg_csat', 'f2')} | {_fmt_prior('avg_csat', 'f2')} | {_wow_cell('avg_csat', 'f2')} | -- | {_status('avg_csat')} |
| Avg Cost/Ticket | {_fmt_val('avg_cost_usd', 'money')} | {_fmt_prior('avg_cost_usd', 'money')} | {_wow_cell('avg_cost_usd', 'money')} | ~${monthly_cost_est:,.0f}/mo | {_status('avg_cost_usd')} |
| CSAT Collection Rate | {_fmt_val('csat_collection_rate')} | {_fmt_prior('csat_collection_rate')} | {_wow_cell('csat_collection_rate')} | -- | {_status('csat_collection_rate')} |

---

#### Data Quality Notes

- Rows processed: {quality_report['total_rows']:,}
- Data completeness: {quality_report['completeness_score']:.1%}
- Fixes applied: {', '.join(f'{k}: {v}' for k, v in cleaning_log.items())}
- Complete weeks used: {', '.join(f'Week {w}' for w in _cw_brief)}

---

*Prepared by: AI-Powered Customer Ops Command Center  |  Groupon Global Customer Operations*
"""

    return data_section


def _generate_html_bytes(markdown_text: str) -> bytes:
    """Convert the brief Markdown into a styled standalone HTML document."""
    import html as _html
    import re as _re

    # Minimal Markdown → HTML conversion (tables, bold, headers, bullets, hr)
    lines = markdown_text.split("\n")
    html_lines: list[str] = []
    in_table = False
    in_ul = False

    for line in lines:
        stripped = line.strip()

        # Horizontal rule
        if stripped in ("---", "***", "___"):
            if in_ul:
                html_lines.append("</ul>")
                in_ul = False
            if in_table:
                html_lines.append("</tbody></table>")
                in_table = False
            html_lines.append("<hr>")
            continue

        # Table rows
        if "|" in stripped and stripped.startswith("|"):
            cells = [c.strip() for c in stripped.strip("|").split("|")]
            # Skip separator rows like |---|---|
            if all(_re.fullmatch(r":?-+:?", c) for c in cells):
                continue
            if not in_table:
                if in_ul:
                    html_lines.append("</ul>")
                    in_ul = False
                html_lines.append('<table><thead><tr>')
                html_lines.extend(f"<th>{_html.escape(c)}</th>" for c in cells)
                html_lines.append("</tr></thead><tbody>")
                in_table = True
            else:
                html_lines.append("<tr>")
                html_lines.extend(f"<td>{_html.escape(c)}</td>" for c in cells)
                html_lines.append("</tr>")
            continue

        if in_table:
            html_lines.append("</tbody></table>")
            in_table = False

        # Headers
        if stripped.startswith("#"):
            level = len(stripped) - len(stripped.lstrip("#"))
            level = min(level, 6)
            text = stripped.lstrip("# ").strip()
            if in_ul:
                html_lines.append("</ul>")
                in_ul = False
            html_lines.append(f"<h{level}>{_html.escape(text)}</h{level}>")
            continue

        # Bullet points
        if stripped.startswith("- ") or stripped.startswith("* "):
            if not in_ul:
                html_lines.append("<ul>")
                in_ul = True
            text = stripped[2:].strip()
            # Bold
            text = _re.sub(r"\*\*(.+?)\*\*", r"<strong>\1</strong>", text)
            html_lines.append(f"<li>{text}</li>")
            continue

        if in_ul:
            html_lines.append("</ul>")
            in_ul = False

        # Empty lines
        if not stripped:
            html_lines.append("")
            continue

        # Regular paragraph with bold support
        text = _re.sub(r"\*\*(.+?)\*\*", r"<strong>\1</strong>", _html.escape(stripped))
        html_lines.append(f"<p>{text}</p>")

    if in_ul:
        html_lines.append("</ul>")
    if in_table:
        html_lines.append("</tbody></table>")

    body = "\n".join(html_lines)

    html_doc = f"""<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="utf-8">
<meta name="viewport" content="width=device-width, initial-scale=1">
<title>Ops Intelligence Brief — Groupon</title>
<style>
  body {{ font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, sans-serif;
         max-width: 900px; margin: 2rem auto; padding: 0 1rem; color: #3d3a2a; line-height: 1.6; }}
  h1,h2,h3 {{ color: #1a1a1a; }}
  h1 {{ border-bottom: 3px solid #53A318; padding-bottom: 0.3em; }}
  table {{ border-collapse: collapse; width: 100%; margin: 1rem 0; }}
  th, td {{ border: 1px solid #ddd; padding: 8px 12px; text-align: left; }}
  th {{ background: #53A318; color: white; }}
  tr:nth-child(even) {{ background: #f9f9f6; }}
  hr {{ border: none; border-top: 2px solid #e0e0e0; margin: 1.5rem 0; }}
  ul {{ padding-left: 1.5rem; }}
  li {{ margin-bottom: 0.3rem; }}
  .footer {{ text-align: center; color: #999; font-size: 0.85em; margin-top: 3rem; }}
</style>
</head>
<body>
{body}
<div class="footer">Prepared by AI-Powered Customer Ops Command Center — Groupon</div>
</body>
</html>"""

    return html_doc.encode("utf-8")


def _generate_word_bytes(markdown_text: str) -> bytes:
    """Convert the brief Markdown into a .docx Word document (in-memory bytes).

    Properly renders Markdown tables as Word tables with headers and borders.
    """
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.oxml.ns import qn
        from docx.enum.table import WD_TABLE_ALIGNMENT
    except ImportError:
        logger.warning("python-docx not installed; Word export unavailable")
        return b""

    doc = Document()

    # Default paragraph style tweaks
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    def _set_cell_border(cell, **kwargs):
        """Set cell borders (top, bottom, left, right)."""
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = tcPr.find(qn("w:tcBorders"))
        if tcBorders is None:
            from lxml import etree
            tcBorders = etree.SubElement(tcPr, qn("w:tcBorders"))
        for edge, val in kwargs.items():
            element = tcBorders.find(qn(f"w:{edge}"))
            if element is None:
                from lxml import etree
                element = etree.SubElement(tcBorders, qn(f"w:{edge}"))
            element.set(qn("w:val"), "single")
            element.set(qn("w:sz"), "4")
            element.set(qn("w:space"), "0")
            element.set(qn("w:color"), "999999")

    # Parse the markdown line-by-line; accumulate table rows
    lines = markdown_text.splitlines()
    i = 0
    while i < len(lines):
        stripped = lines[i].strip()

        # Skip empty lines and horizontal rules
        if not stripped or stripped == "---":
            i += 1
            continue

        # Headings
        if stripped.startswith("####"):
            doc.add_heading(stripped.lstrip("# ").strip(), level=2)
            i += 1
            continue
        if stripped.startswith("###"):
            doc.add_heading(stripped.lstrip("# ").strip(), level=1)
            i += 1
            continue
        if stripped.startswith("##"):
            doc.add_heading(stripped.lstrip("# ").strip(), level=1)
            i += 1
            continue

        # Detect Markdown table block
        if stripped.startswith("|") and "|" in stripped[1:]:
            table_rows: list[list[str]] = []
            while i < len(lines):
                row_text = lines[i].strip()
                if not row_text.startswith("|"):
                    break
                # Skip separator rows (|---|---|)
                cells = [c.strip() for c in row_text.strip("|").split("|")]
                if all(set(c) <= {"-", ":", " "} for c in cells):
                    i += 1
                    continue
                table_rows.append(cells)
                i += 1

            if table_rows:
                n_cols = max(len(r) for r in table_rows)
                table = doc.add_table(rows=len(table_rows), cols=n_cols)
                table.alignment = WD_TABLE_ALIGNMENT.CENTER

                for row_idx, row_data in enumerate(table_rows):
                    for col_idx, cell_text in enumerate(row_data):
                        if col_idx >= n_cols:
                            break
                        cell = table.cell(row_idx, col_idx)
                        cell.text = cell_text.replace("**", "")
                        # Style header row
                        if row_idx == 0:
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    run.bold = True
                                    run.font.size = Pt(10)
                        else:
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    run.font.size = Pt(10)
                        _set_cell_border(
                            cell, top="single", bottom="single",
                            left="single", right="single",
                        )
                doc.add_paragraph()  # spacing after table
            continue

        # Bullet points
        if stripped.startswith("- "):
            text = stripped.lstrip("- ").replace("**", "")
            doc.add_paragraph(text, style="List Bullet")
            i += 1
            continue

        # Footer / signature
        if stripped.startswith("*Prepared"):
            p = doc.add_paragraph()
            run = p.add_run(stripped.strip("*"))
            run.italic = True
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
            i += 1
            continue

        # Normal paragraph — strip bold markers
        doc.add_paragraph(stripped.replace("**", ""))
        i += 1

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _generate_pdf_bytes(
    analytics: dict,
    nlp_results: dict,
    quality_report: dict,
    cleaning_log: dict,
    *,
    markdown_text: str = "",
) -> bytes:
    """Generate a PDF report from the weekly brief using fpdf2."""
    try:
        from fpdf import FPDF
    except ImportError:
        logger.warning("fpdf2 not installed; PDF export unavailable")
        return b""

    # fpdf2's built-in Helvetica only supports latin-1.
    # Replace common Unicode chars with safe ASCII equivalents.
    _UNICODE_MAP = {
        "\u2014": "--",   # em dash
        "\u2013": "-",    # en dash
        "\u2018": "'",    # left single quote
        "\u2019": "'",    # right single quote
        "\u201c": '"',    # left double quote
        "\u201d": '"',    # right double quote
        "\u2026": "...",  # ellipsis
        "\u2022": "-",    # bullet
        "\u2190": "<-",   # left arrow
        "\u2192": "->",   # right arrow
        "\u2265": ">=",   # >=
        "\u2264": "<=",   # <=
        "\u00a0": " ",    # non-breaking space
        "\u200b": "",     # zero-width space
    }

    def _sanitize(text: str) -> str:
        for uc, repl in _UNICODE_MAP.items():
            text = text.replace(uc, repl)
        # Catch any remaining non-latin-1 chars
        return text.encode("latin-1", errors="replace").decode("latin-1")

    class _BriefPDF(FPDF):
        """Custom PDF with header/footer branding."""

        def header(self):
            self.set_font("Helvetica", "B", 10)
            self.set_text_color(83, 163, 24)  # Groupon green
            self.cell(0, 8, "Groupon Ops Intelligence Brief", new_x="LMARGIN", new_y="NEXT", align="L")
            self.set_draw_color(83, 163, 24)
            self.line(10, self.get_y(), self.w - 10, self.get_y())
            self.ln(4)

        def footer(self):
            self.set_y(-15)
            self.set_font("Helvetica", "I", 8)
            self.set_text_color(128, 128, 128)
            self.cell(0, 10, f"Page {self.page_no()}/{{nb}}", align="C")

    pdf = _BriefPDF()
    pdf.alias_nb_pages()
    pdf.set_auto_page_break(auto=True, margin=20)
    pdf.add_page(orientation="L")  # landscape for wide KPI table

    # Render the markdown text as formatted PDF content
    text = _sanitize(markdown_text or "No brief content available.")

    lines = text.splitlines()
    in_table = False
    table_rows: list[list[str]] = []

    def _flush_table():
        """Render accumulated table rows."""
        nonlocal table_rows
        if not table_rows:
            return
        n_cols = max(len(r) for r in table_rows)
        if n_cols == 0:
            table_rows = []
            return
        page_w = pdf.w - pdf.l_margin - pdf.r_margin  # works for any orientation
        col_w = page_w / n_cols

        for row_idx, row in enumerate(table_rows):
            for col_idx in range(n_cols):
                cell_text = row[col_idx] if col_idx < len(row) else ""
                cell_text = cell_text.replace("**", "")
                if row_idx == 0:
                    pdf.set_font("Helvetica", "B", 9)
                    pdf.set_fill_color(83, 163, 24)
                    pdf.set_text_color(255, 255, 255)
                else:
                    pdf.set_font("Helvetica", "", 9)
                    pdf.set_fill_color(245, 245, 245) if row_idx % 2 == 0 else pdf.set_fill_color(255, 255, 255)
                    pdf.set_text_color(0, 0, 0)

                pdf.cell(col_w, 6, cell_text[:50], border=1, fill=True)
            pdf.ln()
        pdf.set_text_color(0, 0, 0)
        pdf.set_x(pdf.l_margin)  # guarantee X is reset after table
        pdf.ln(3)
        table_rows = []

    for line in lines:
        stripped = line.strip()

        # Skip empty / horizontal rules
        if not stripped:
            if in_table:
                _flush_table()
                in_table = False
            pdf.ln(2)
            continue
        if stripped == "---":
            if in_table:
                _flush_table()
                in_table = False
            pdf.ln(3)
            continue

        # Table rows
        if stripped.startswith("|") and "|" in stripped[1:]:
            cells = [c.strip() for c in stripped.strip("|").split("|")]
            # Skip separator rows
            if all(set(c) <= {"-", ":", " "} for c in cells):
                continue
            in_table = True
            table_rows.append(cells)
            continue

        # Flush any pending table
        if in_table:
            _flush_table()
            in_table = False

        # Headings
        if stripped.startswith("####"):
            pdf.set_font("Helvetica", "B", 12)
            pdf.set_text_color(51, 51, 51)
            pdf.cell(0, 8, stripped.lstrip("# ").strip(), new_x="LMARGIN", new_y="NEXT")
            pdf.ln(2)
        elif stripped.startswith("###") or stripped.startswith("##"):
            pdf.set_font("Helvetica", "B", 14)
            pdf.set_text_color(83, 163, 24)
            pdf.cell(0, 10, stripped.lstrip("# ").strip(), new_x="LMARGIN", new_y="NEXT")
            pdf.ln(3)
        elif stripped.startswith("- "):
            pdf.set_font("Helvetica", "", 10)
            pdf.set_text_color(0, 0, 0)
            bullet_text = stripped[2:].replace("**", "").replace("*", "")
            # Force X to left margin before multi_cell to guarantee
            # enough horizontal space.
            pdf.set_x(pdf.l_margin)
            pdf.multi_cell(0, 5, f"  -  {bullet_text}")
        elif stripped.startswith("*Prepared"):
            pdf.ln(5)
            pdf.set_font("Helvetica", "I", 8)
            pdf.set_text_color(128, 128, 128)
            pdf.cell(0, 5, stripped.strip("*"), new_x="LMARGIN", new_y="NEXT")
        else:
            pdf.set_font("Helvetica", "", 10)
            pdf.set_text_color(0, 0, 0)
            clean = stripped.replace("**", "").replace("*", "")
            pdf.set_x(pdf.l_margin)
            pdf.multi_cell(0, 5, clean)
            pdf.ln(1)

    # Flush any trailing table
    if in_table:
        _flush_table()

    try:
        buf = io.BytesIO()
        pdf.output(buf)
        return buf.getvalue()
    except Exception as exc:
        logger.error("PDF output failed: %s", exc)
        return b""


def _compute_watch_list(analytics: dict, nlp_results: dict) -> list[str]:
    """Generate a deterministic watch list of emerging patterns.

    Flags metrics that are worsening >5% WoW, anomaly thresholds,
    and sentiment shifts.
    """
    items: list[str] = []
    wow = analytics.get("wow", {})
    deltas = wow.get("deltas", {})

    # Flag any KPI worsening >5% WoW
    metric_labels = {
        "resolution_rate": "Resolution rate",
        "escalation_rate": "Escalation rate",
        "abandonment_rate": "Abandonment rate",
        "avg_first_response_min": "Avg first response time",
        "avg_csat": "Avg CSAT",
        "avg_cost_usd": "Avg cost per ticket",
    }
    for key, label in metric_labels.items():
        d = deltas.get(key, {})
        pct = d.get("pct_change", 0)
        direction = d.get("direction", "stable")
        if direction == "worsening" and abs(pct) > 5:
            items.append(f"**{label}** worsening {abs(pct):.1f}% week-over-week")

    # Flag high chatbot escalation
    chatbot = analytics.get("chatbot", {})
    esc_rate = chatbot.get("overall_escalation_rate", 0)
    if esc_rate > 0.20:
        items.append(
            f"**Chatbot escalation rate** at {esc_rate:.1%} — above 20% threshold"
        )

    # Flag frustration rate
    frust_rate = nlp_results.get("frustration_rate", 0)
    if frust_rate > 0.15:
        items.append(
            f"**Customer frustration rate** at {frust_rate:.1%} — "
            "consider targeted improvements in top frustration categories"
        )

    # Flag volume trend (growing?)
    total_delta = deltas.get("total_tickets", {})
    if total_delta.get("direction") == "worsening" or total_delta.get("pct_change", 0) > 5:
        items.append(
            "**Ticket volume** trending upward — monitor for capacity planning"
        )

    return items


def render_weekly_brief_tab(
    analytics: dict,
    nlp_results: dict,
    quality_report: dict,
    cleaning_log: dict,
    agent_result: dict | None = None,
) -> None:
    """Render the weekly brief -- agent narrative + deterministic KPI reference."""

    if agent_result is None:
        agent_result = st.session_state.get("agent_result")

    # Week context
    wow = analytics.get("wow", {})
    cur_wk = wow.get("current_week")
    pri_wk = wow.get("prior_week")
    all_weeks = sorted(analytics.get("complete_weeks", COMPLETE_WEEKS))
    date_ranges = analytics.get("week_date_ranges", {})
    cur_dates = date_ranges.get(cur_wk, "")
    pri_dates = date_ranges.get(pri_wk, "")
    first_dates = date_ranges.get(min(all_weeks) if all_weeks else 0, "")
    last_dates = date_ranges.get(max(all_weeks) if all_weeks else 0, "")

    if len(all_weeks) > 2:
        st.info(
            f"Data spans Weeks {min(all_weeks)}-{max(all_weeks)} "
            f"({first_dates} to {last_dates}, {len(all_weeks)} complete weeks). "
            f"This brief focuses on the most recent two: "
            f"**Week {cur_wk}** ({cur_dates}) vs **Week {pri_wk}** ({pri_dates}).",
            icon="📅",
        )

    # -- 1. Agent Executive Brief (primary content) --
    ai_brief = agent_result.get("report_markdown") if agent_result else None
    if ai_brief:
        st.markdown(ai_brief)
    else:
        st.warning(
            "AI executive brief unavailable -- the agent pipeline may have "
            "encountered an error.  The deterministic KPI reference below "
            "is always available."
        )

    st.divider()

    # -- 2. Deterministic KPI Reference (grounded numbers) --
    kpi_markdown = _generate_markdown_brief(
        analytics, nlp_results, quality_report, cleaning_log,
    )
    with st.expander(
        "Verified KPI Reference (deterministic)",
        expanded=not bool(ai_brief),
    ):
        st.caption(
            "These numbers are computed deterministically from the raw data "
            "and serve as the authoritative reference."
        )
        st.markdown(kpi_markdown)

    st.divider()

    # -- 3. Watch List (emerging patterns) --
    st.subheader("⚠️ Watch List — Emerging Patterns")
    watch_items = _compute_watch_list(analytics, nlp_results)
    if watch_items:
        for item in watch_items:
            st.markdown(f"- ⚠️ {item}")
    else:
        st.caption("No emerging patterns flagged this period.")

    st.divider()

    # -- 4. Export (agent brief + KPI reference combined) --
    st.subheader("Export Report")
    export_md = ""
    if ai_brief:
        export_md += ai_brief + "\n\n---\n\n"
    export_md += kpi_markdown

    col_pdf, col_wd, col_html, col_md = st.columns(4)

    with col_pdf:
        pdf_bytes = _generate_pdf_bytes(
            analytics, nlp_results, quality_report, cleaning_log,
            markdown_text=export_md,
        )
        if pdf_bytes:
            st.download_button(
                label="PDF",
                data=pdf_bytes,
                file_name="ops_brief.pdf",
                mime="application/pdf",
                key="dl_pdf",
            )
        else:
            st.button("PDF (unavailable)", disabled=True)

    with col_wd:
        docx_bytes = _generate_word_bytes(export_md)
        if docx_bytes:
            st.download_button(
                label="Word",
                data=docx_bytes,
                file_name="ops_brief.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                key="dl_word",
            )
        else:
            st.button("Word (unavailable)", disabled=True)

    with col_html:
        html_bytes = _generate_html_bytes(export_md)
        st.download_button(
            label="HTML",
            data=html_bytes,
            file_name="ops_brief.html",
            mime="text/html",
            key="dl_html",
        )

    with col_md:
        st.download_button(
            label="Markdown",
            data=export_md,
            file_name="ops_brief.md",
            mime="text/markdown",
            key="dl_md",
        )

    # Agent execution log
    if agent_result:
        with st.expander("Agent Execution Log"):
            log_entries = agent_result.get("execution_log", [])
            for entry in log_entries:
                st.write(f"- {entry}")
            if not log_entries:
                st.write("No execution logs found.")


